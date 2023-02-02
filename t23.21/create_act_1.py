import openpyxl
import docx
from docx.shared import Pt, Cm


# зчитуємо необхідні дані
# передаємономер рахунку та файл де зберігаються дані
def create_invoice(act_no, datafile):
    wb = openpyxl.load_workbook(datafile)  # відкриваємо файл

    # відкриваємо робочий аркуш рахунку
    ws = wb["acts"]
    # шукаємо у переліку рахунок з відповідним номером
    for row in ws:
        if row[1].value == act_no:
            # якщо дані однакові, зберігаємо
            act_id, act_no, act_date, act_sum, fields_s_id = [c.value for c in row]
            #  print(invoice_id, invoice_no, invoice_date, customer_id)
            break
    else:
        raise RuntimeError

    # тепер ми маємо індентифікатор клієнта і за ним можемо шукати клієнтів
    ws = wb["fields"]
    for row in ws:
        if row[0].value == fields_s_id:  # ідентифікатор - перший елемент у рядку
            fields_s_id, fields_name, fields_address, responsible_person, manager = [c.value for c in row]
            # print(customer_id, customer_name, customer_address)
            break
    else:
        raise RuntimeError

    works = {}  # ключ - ідентифікатор продукту, а значення - словник з інформацією про даний продукт
    # ідентифікатор також присутній у таблиці items
    # знайдемо всі ідентифікатори даного замовлення, які його стосуються
    ws = wb["items"]
    for row in ws:
        if row[1].value == act_id:  # перший елемент співпадає з ідентифікатором рахунку
            works_id, act_id = [c.value for c in row]
            # print(invoice_id, product_id, quantity)
            print(works_id, act_id)
            works[works_id] = {}  # записуємо інформацію про продукт у словник

    # з'ясуємо інформацію про дані продукту
    ws = wb["works"]
    for row in ws:
        if row[0].value in works:  # ми повинні записати для кожного продукту
            # чи є продукт у словнику продуктів за даним замовленням
            works_id, works_name = [c.value for c in row]
            works[works_id]["name"] = works_name  # додаємо інформацію у внутрішній словник
            works[works_id]["price"] = act_sum
            works[works_id]["responsible person"] = responsible_person
            works[works_id]["manager"] = manager

    print(works)

    # формуємо документ
    doc = docx.Document()
    # надаємо стиль всьому документу
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    # print(style.font.size)

    # створимо таблицю (1 рядок і 2 стовпчики), перший елемент вирівнювати по лівому краю, а другий - по правому
    table = doc.add_table(1, 2)
    paragraph = table.cell(0, 0).add_paragraph(f"Акт № {act_no}")
    paragraph.alignment = 0
    paragraph = table.cell(0, 1).add_paragraph(f"Дата {act_date}")
    paragraph.alignment = 2

    # запишемо інформацію про покупця
    doc.add_paragraph(f"Даний Акт засвідчує, що Виконавцем на майданчику {fields_name} за адресою {fields_address}"
                      f" були виконані такі роботи:")

    headers = ("№", "Назва роботи")
    table = doc.add_table(1, len(headers), "Table Grid")
    # 1 рядок, к-ть стовпчиків = к-ті заголовків, "Table Grid" - малюємо границі таблиці
    row = table.rows[0]  # отримуємо рядок зі списку рядків
    for cell, header in zip(row.cells, headers):  # одночасно ітеруємо і по клітинках і по заголовках
        cell.text = header

    total = 0
    for i, work in enumerate(works.values(), 1):
        full_price = int(work["price"])
        responsible_person = work["responsible person"]
        manager = work["manager"]
        total += full_price
        values = (
            i,
            work["name"]
        )
        row = table.add_row()  # додаємо рядок
        for cell, value in zip(row.cells, values):
            cell.text = str(value)

    # змінимо розмір клітин
    widths = (Cm(1.0), Cm(20))  # розміри колонок, остання забере все, що залишиться
    for column, width in zip(table.columns, widths):
        for cell in column.cells:
            cell.width = width

    paragraph = doc.add_paragraph(f"Сума виконаних робіт складає {total} грн.")  # вирівнюємо по правому краю
    paragraph.alignment = 0
    paragraph.space_after = Pt(15)
    table = doc.add_table(2, 3)
    paragraph = table.cell(0, 0).add_paragraph(f"Від Замовника")
    paragraph.alignment = 0
    paragraph = table.cell(1, 0).add_paragraph(responsible_person)
    paragraph.alignment = 0
    paragraph = table.cell(0, 2).add_paragraph(f"Від Виконавця")
    paragraph.alignment = 2
    paragraph = table.cell(1, 2).add_paragraph(manager)
    paragraph.alignment = 2
    doc.save(f"new_act.docx")


if __name__ == "__main__":
    create_invoice(34, "data.xlsx")
